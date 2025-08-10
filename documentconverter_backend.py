# documentconverter_backend.py
from PIL import Image
from pdf2image import convert_from_path
from docx import Document
import pytesseract
import pdfplumber
import io
import cv2
import os
import numpy as np
import docx

# 1. Convert multiple images to one PDF
def convert_images_to_pdf(image_files, resize_dims=None, output_path="output_images.pdf"):
    images = []
    for file in image_files:
        img = Image.open(file).convert("RGB")
        if resize_dims:
            img = img.resize(resize_dims)
        images.append(img)
    if images:
        images[0].save(output_path, save_all=True, append_images=images[1:])
    return output_path

# 2. Convert PDF to images
def convert_pdf_to_images(pdf_file, output_dir="pdf_pages"):
    os.makedirs(output_dir, exist_ok=True)
    images = convert_from_path(pdf_file)
    saved_paths = []
    for i, img in enumerate(images):
        img_path = os.path.join(output_dir, f"page_{i+1}.jpg")
        img.save(img_path, "JPEG")
        saved_paths.append(img_path)
    return saved_paths

# 3. OCR: Image to editable Word
def ocr_image_to_word(image_file, output_path="ocr_result.docx"):
    img = Image.open(image_file)
    text = pytesseract.image_to_string(img)
    doc = Document()
    doc.add_paragraph(text)
    doc.save(output_path)
    return output_path

# 4. Convert image to grayscale (with optional resize)
def grayscale_image(image_file, resize_dims=None, output_path="grayscale.jpg"):
    img = cv2.imdecode(np.frombuffer(image_file.read(), np.uint8), cv2.IMREAD_COLOR)
    if resize_dims:
        img = cv2.resize(img, resize_dims)
    gray = cv2.cvtColor(img, cv2.COLOR_BGR2GRAY)
    cv2.imwrite(output_path, gray)
    return output_path

# 5. Convert PDF to Word
def pdf_to_word(pdf_file, output_path="converted_from_pdf.docx"):
    doc = Document()
    with pdfplumber.open(pdf_file) as pdf:
        for page in pdf.pages:
            text = page.extract_text()
            if text:
                doc.add_paragraph(text)
    doc.save(output_path)
    return output_path

# 6. Convert PDF to Text
def pdf_to_text(pdf_file, output_path="output_pdf_text.txt"):
    text_output = ""
    with pdfplumber.open(pdf_file) as pdf:
        for page in pdf.pages:
            page_text = page.extract_text()
            if page_text:
                text_output += page_text + "\n"
    with open(output_path, "w", encoding="utf-8") as f:
        f.write(text_output)
    return output_path

# 7. Convert Word to Text
def word_to_text(docx_file, output_path="output_word_text.txt"):
    doc = docx.Document(docx_file)
    text_output = "\n".join([para.text for para in doc.paragraphs])
    with open(output_path, "w", encoding="utf-8") as f:
        f.write(text_output)
    return output_path
